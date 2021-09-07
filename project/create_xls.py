import os
import sys
import docx
import pandas as pd

"""
- Import Questionnaire from MS Word document
- Extract text from each row, return as dictionary
    - If text contains Section text, this is begin_group type
    - Column 0 - if text 'Q' this is a question. Strangely, docx is not importing the question numbers?
        - Autonumber each question. May not match the numbering in the questionnaire
    - Column 1 - this is question label
    - Column 2(quest_choices) - If not empty, assume this is a select_one or select_multiple question
    - Column 3 - if not empty, these are the choice values of the question
- Survey dictionary - create 'survey' and 'choices' sheets from this dictionary
    - Survey sheet - has the following fields
        - type - if quest_choices not empty, type should be select_one. Otherwise, integer or text. 
        Each select_one needs a list_name
        - name - unique name for each question. For now add Q to question number e.g., Q1
        - label - question_label
    - Choices sheet - has the following fields:
        - list_name - this is repeating field depending on the number of choice values
        - name - unique number/text for each choice value within the list_name
        - label - quest_choices
        
"""

# Update these values after running the inspect_table function.
LABEL_CELL_NUM = 1
CHOICE_CELL_NUM = 2
CHOICE_LABEL_NUM = 3


def inspect_table(doc):
    """ Inspect tables in Microsoft Word re cell numbers with question_text, choices and choice labels

    :param doc: Microsoft Word document containing tables
    :return:
    """
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text == 'Q':
                for cell_num in range(len(row.cells)):
                    print(f"Cell Number: {cell_num},"
                          f"Cell text: {row.cells[cell_num].text}\n")


def survey_dict(doc):
    """ Produce dictionary consisting of survey questions

    :param doc: Microsoft Word document containing tables
    :return: Dic of survey questions - key, question type and label
    """
    survey = {}
    num = 1
    for table in doc.tables:
        # print(len(table.rows))
        for row in table.rows:
            # print(len(row.cells))
            if row.cells[0].text == 'Q':
                quest_id = num
                quest_label = row.cells[LABEL_CELL_NUM].text
                # print(quest_label)
                quest_choices = row.cells[CHOICE_CELL_NUM].text.split('\n')
                # print(f"Choices for this question{quest_choices}")
                quest_choice_labels = row.cells[CHOICE_LABEL_NUM].text.split('\n')
                survey[quest_id] = quest_label, quest_choices, quest_choice_labels
                num += 1
    return survey


def survey_choices(survey):
    """Generate survey and choices sheets

    :param survey: Dictionary containing the survey
    :return: survey sheet(Dictionary) and choices_sheet(List)
    """
    survey_sheet = {}
    choices_sheet = []
    for key, quest in survey.items():
        # print(quest_id, quest)
        quest_code = 'GSQ'  # change this depending on the questionnaire
        quest_id = quest_code + str(key)
        quest_label = quest[0]
        if len(quest[1]) <= 1:
            quest_type = 'text'
        else:
            list_name = quest_id
            choices = quest[1]
            choice_values = quest[2]
            # Check if the number of choices match number of choice values
            if len(choices) != len(choice_values):
                choice_values = list(range(len(choices)))
                for i in range(len(choices)):
                    choice_values[i] = 900 + i  # populate choice_values with dummy variables
            # Choices sheet
            quest_type = 'select_one ' + list_name
            for i in range(len(choices)):
                choices_sheet.append((list_name, choices[i], choice_values[i]))
        survey_sheet[key] = quest_id, quest_type, quest_label
    return survey_sheet, choices_sheet


def main():
    doc = docx.Document("data/survey_ke_2.docx")

    survey_sheet, choices_sheet = survey_choices(survey_dict(doc))

    unwanted_xters = ['…', '…', '_']

    survey_cols = ['name', 'type', 'label']
    choices_cols = ['list_name', 'label', 'name']

    # print(survey_sheet)
    survey_df = pd.DataFrame.from_dict(survey_sheet, orient='index', columns=survey_cols)
    choices_df = pd.DataFrame(choices_sheet, columns=choices_cols)
    # print(choices_df)

    with pd.ExcelWriter("data/survey_ke.xlsx") as writer:
        survey_df.to_excel(writer, sheet_name='survey')
        choices_df.to_excel(writer, sheet_name='choices')

    # Inspect questions and choices - may be redundant since Panda does that, ensuring all data is within set columns
    with open("data/survey_sheet.txt", "w") as survey_file:
        for quest_id, quest in survey_sheet.items():
            survey_file.write(str(quest_id) + " - " + quest[0] + ": " + quest[1] + "\n")

    with open("data/choices_sheet.txt", "w") as choices_file:
        for choice in choices_sheet:
            list_name = choice[0]
            label = choice[1]
            num = choice[2]
            choices_file.write(list_name + " " + label + " " + str(num) + "\n")
    return os.EX_OK


if __name__ == "__main__":
    sys.exit(main())
